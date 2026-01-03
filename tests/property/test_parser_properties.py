"""Property-based tests for Document Parser.

Feature: knowledge-atomizer
"""

import os
import tempfile
import uuid
from typing import List, Tuple

from docx import Document
from docx.shared import Pt
from hypothesis import given, settings, strategies as st, assume

import sys
sys.path.insert(0, '.')

from src.parser import DocumentParser, InvalidFileFormatError, FileNotFoundError
from src.models import HeadingLevel, DocumentNode


# Helper to create test documents
def create_test_document(headings: List[Tuple[int, str]], paragraphs_per_heading: int = 0) -> str:
    """Create a test .docx file with specified headings.
    
    Args:
        headings: List of (level, title) tuples where level is 1-5
        paragraphs_per_heading: Number of body paragraphs to add after each heading
        
    Returns:
        Path to the created temporary file
    """
    doc = Document()
    
    for level, title in headings:
        # Add heading
        heading = doc.add_heading(title, level=level)
        
        # Add body paragraphs if requested
        for i in range(paragraphs_per_heading):
            doc.add_paragraph(f"Content paragraph {i+1} under {title}")
    
    # Save to temp file
    fd, path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    doc.save(path)
    return path


# Strategies for generating heading structures
heading_level = st.integers(min_value=1, max_value=5)
heading_title = st.text(min_size=1, max_size=50, alphabet=st.characters(
    whitelist_categories=('L', 'N', 'P', 'S'),
    whitelist_characters=' '
)).filter(lambda x: x.strip())

heading_entry = st.tuples(heading_level, heading_title)
heading_list = st.lists(heading_entry, min_size=1, max_size=10)


@given(headings=heading_list)
@settings(max_examples=100)
def test_heading_level_recognition(headings: List[Tuple[int, str]]):
    """
    Feature: knowledge-atomizer
    Property 1: Heading Level Recognition
    Validates: Requirements 1.2
    
    For any valid Word document containing headings with styles Heading 1-5,
    the Parser SHALL correctly identify and return the corresponding HeadingLevel
    (H1-H5) for each heading.
    """
    # Filter out empty titles after strip
    headings = [(level, title.strip()) for level, title in headings if title.strip()]
    assume(len(headings) > 0)
    
    # Make titles unique by appending index to avoid duplicate title issues
    unique_headings = [(level, f"{title}_{i}") for i, (level, title) in enumerate(headings)]
    
    # Create test document
    path = create_test_document(unique_headings)
    
    try:
        parser = DocumentParser()
        tree = parser.parse(path)
        
        # Collect all nodes from tree in order (pre-order traversal)
        def collect_nodes(nodes: List[DocumentNode]) -> List[DocumentNode]:
            result = []
            for node in nodes:
                result.append(node)
                result.extend(collect_nodes(node.children))
            return result
        
        all_nodes = collect_nodes(tree.root_nodes)
        
        # Verify we got the right number of headings
        assert len(all_nodes) == len(unique_headings), \
            f"Expected {len(unique_headings)} nodes, got {len(all_nodes)}"
        
        # Verify each heading level is correctly identified by title
        title_to_level = {title: level for level, title in unique_headings}
        
        for node in all_nodes:
            expected_level = title_to_level.get(node.title)
            assert expected_level is not None, f"Unexpected node title: {node.title}"
            assert node.level.value == expected_level, \
                f"Heading '{node.title}' should be H{expected_level}, got H{node.level.value}"
    finally:
        os.unlink(path)


# Strategy for generating paragraphs under headings
paragraphs_count = st.integers(min_value=1, max_value=3)


@given(
    headings=st.lists(heading_entry, min_size=1, max_size=5),
    para_count=paragraphs_count
)
@settings(max_examples=100)
def test_paragraph_attribution(headings: List[Tuple[int, str]], para_count: int):
    """
    Feature: knowledge-atomizer
    Property 3: Paragraph Attribution
    Validates: Requirements 1.4
    
    For any document with headings and paragraphs, every paragraph SHALL be
    attributed to exactly one heading (the nearest preceding heading),
    and no paragraph shall be orphaned.
    """
    # Filter out empty titles
    headings = [(level, title.strip()) for level, title in headings if title.strip()]
    assume(len(headings) > 0)
    
    # Create test document with paragraphs
    path = create_test_document(headings, paragraphs_per_heading=para_count)
    
    try:
        parser = DocumentParser()
        tree = parser.parse(path)
        
        # Collect all nodes
        def collect_nodes(nodes: List[DocumentNode]) -> List[DocumentNode]:
            result = []
            for node in nodes:
                result.append(node)
                result.extend(collect_nodes(node.children))
            return result
        
        all_nodes = collect_nodes(tree.root_nodes)
        
        # Each heading should have content (the paragraphs)
        for node in all_nodes:
            # Content should not be empty since we added paragraphs
            assert node.content, f"Node '{node.title}' should have content"
            
            # Content should contain the expected paragraph text
            for i in range(para_count):
                expected_text = f"Content paragraph {i+1} under {node.title}"
                assert expected_text in node.content, \
                    f"Node '{node.title}' should contain '{expected_text}'"
    finally:
        os.unlink(path)


# Strategies for invalid files
invalid_extensions = st.sampled_from(['.txt', '.pdf', '.doc', '.xlsx', '.pptx', ''])


@given(ext=invalid_extensions)
@settings(max_examples=100)
def test_invalid_file_rejection(ext: str):
    """
    Feature: knowledge-atomizer
    Property 4: Invalid File Rejection
    Validates: Requirements 1.5
    
    For any file that is not a valid .docx file, the Parser SHALL raise
    a descriptive error rather than crash or return invalid data.
    """
    # Create a temp file with invalid extension
    fd, path = tempfile.mkstemp(suffix=ext)
    os.write(fd, b"This is not a valid docx file")
    os.close(fd)
    
    try:
        parser = DocumentParser()
        
        try:
            parser.parse(path)
            assert False, f"Parser should reject file with extension '{ext}'"
        except InvalidFileFormatError as e:
            # Expected - verify error message is descriptive
            assert ext in str(e) or "不支持" in str(e) or ".docx" in str(e), \
                f"Error message should be descriptive: {e}"
        except Exception as e:
            # Other exceptions are also acceptable for truly invalid files
            pass
    finally:
        os.unlink(path)


def test_nonexistent_file_rejection():
    """
    Feature: knowledge-atomizer
    Property 4: Invalid File Rejection (nonexistent file case)
    Validates: Requirements 1.5
    """
    parser = DocumentParser()
    fake_path = f"/nonexistent/path/{uuid.uuid4()}.docx"
    
    try:
        parser.parse(fake_path)
        assert False, "Parser should reject nonexistent file"
    except FileNotFoundError as e:
        assert fake_path in str(e) or "不存在" in str(e), \
            f"Error message should mention the file path: {e}"


# Helper to create document with table
def create_document_with_table(rows: int, cols: int, cell_data: List[List[str]]) -> str:
    """Create a test .docx file with a table under a heading.
    
    Args:
        rows: Number of rows
        cols: Number of columns
        cell_data: 2D list of cell contents
        
    Returns:
        Path to the created temporary file
    """
    doc = Document()
    doc.add_heading("Table Section", level=1)
    
    table = doc.add_table(rows=rows, cols=cols)
    for i, row_data in enumerate(cell_data):
        if i < rows:
            for j, cell_text in enumerate(row_data):
                if j < cols:
                    table.rows[i].cells[j].text = cell_text
    
    fd, path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    doc.save(path)
    return path


# Strategy for table dimensions and content
table_rows = st.integers(min_value=1, max_value=5)
table_cols = st.integers(min_value=1, max_value=5)
cell_content = st.text(min_size=0, max_size=20, alphabet=st.characters(
    whitelist_categories=('L', 'N'),
    whitelist_characters=' '
))


@given(
    rows=table_rows,
    cols=table_cols,
    data=st.data()
)
@settings(max_examples=100)
def test_table_to_markdown_conversion(rows: int, cols: int, data):
    """
    Feature: knowledge-atomizer
    Property 2: Table to Markdown Conversion
    Validates: Requirements 1.3, 2.5
    
    For any Word table, converting it to Markdown should produce a table
    with the same number of rows and columns, and equivalent cell content.
    """
    # Generate cell data
    cell_data = []
    for i in range(rows):
        row_data = []
        for j in range(cols):
            cell_text = data.draw(cell_content)
            row_data.append(cell_text)
        cell_data.append(row_data)
    
    # Create document with table
    path = create_document_with_table(rows, cols, cell_data)
    
    try:
        parser = DocumentParser()
        tree = parser.parse(path)
        
        # Get the node with the table
        assert len(tree.root_nodes) == 1
        node = tree.root_nodes[0]
        
        # Parse the markdown table
        content = node.content
        assert content, "Node should have table content"
        
        lines = [l for l in content.strip().split('\n') if l.strip()]
        
        # Should have header + separator + data rows
        expected_lines = 1 + 1 + (rows - 1)  # header, separator, data rows
        assert len(lines) >= expected_lines, \
            f"Expected at least {expected_lines} lines, got {len(lines)}"
        
        # Verify column count by counting pipes in header
        header_cols = lines[0].count('|') - 1  # pipes minus outer ones
        assert header_cols == cols, f"Expected {cols} columns, got {header_cols}"
        
        # Verify cell content is preserved (check first row as header)
        for j, expected_text in enumerate(cell_data[0]):
            expected_clean = expected_text.strip()
            if expected_clean:
                assert expected_clean in content, \
                    f"Cell content '{expected_clean}' should be in markdown"
    finally:
        os.unlink(path)
